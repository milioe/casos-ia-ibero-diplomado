import os
import json
from typing import Dict, List, Any, Optional
from pydantic import BaseModel, Field
from openai import OpenAI
from dotenv import load_dotenv
import docx
import io
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Load environment variables
load_dotenv()

class DialogueInfo(BaseModel):
    speakers: List[str] = Field(description="Lista de participantes en la conversación")
    dialogue: List[Dict[str, str]] = Field(description="Lista de intervenciones con el formato {speaker: nombre, text: texto}")

def process_transcript(text: str) -> DialogueInfo:
    """
    Process an unstructured transcript using OpenAI to organize speakers and dialogue.
    
    Args:
        text (str): The unstructured transcript text
        
    Returns:
        DialogueInfo: Structured dialogue information
    """
    try:
        # Initialize OpenAI client
        client = OpenAI(
            api_key=os.getenv("OPENAI_API_KEY"),
            organization=os.getenv("OPENAI_ORG_ID")
        )

        # Create the system message with instructions
        system_message = """Eres un experto en procesar transcripciones y organizar diálogos.
        
        Debes analizar el texto proporcionado y extraer:
        1. Lista de todos los participantes/speakers en la conversación, incluyendo a la persona agraviada.
        2. Diálogo organizado cronológicamente, identificando quién dice qué
        
        Reglas:
        - Identifica correctamente quién está hablando en cada momento
        - Mantén el orden cronológico del diálogo
        - Agrupa el texto de cada intervención de manera coherente
        - Si un speaker no está explícitamente nombrado, úsalo como "Speaker Unknown"
        - Limpia y organiza el texto manteniendo su significado original
        - Siempre habrá una persona agraviada, la cual será nombrada como "Persona Agraviada" en la lista de speakers y en el diálogo
        - Siempre habrá uno o mas entrevistadores, los cuales serán nombrados como "Speaker A" en la lista de speakers y en el diálogo
        
        El formato de salida debe ser un JSON con la siguiente estructura:
        {
            "speakers": ["Lista de speakers, incluyendo 'Persona Agraviada' si corresponde"],
            "dialogue": [
                {
                    "speaker": "Nombre del speaker",
                    "text": "Texto de la intervención"
                }
            ]
        }
        """

        response = client.chat.completions.create(
            model="gpt-4.1",
            messages=[
                {
                    "role": "system",
                    "content": system_message
                },
                {
                    "role": "user",
                    "content": f"Analiza y estructura el siguiente texto:\n\n{text}"
                }
            ],
            response_format={"type": "json_object"},
            temperature=0.1
        )

        # Parse the response into our Pydantic model
        result = response.choices[0].message.content
        dialogue_info = DialogueInfo.model_validate_json(result)
        
        return dialogue_info

    except Exception as e:
        raise Exception(f"Error processing transcript with OpenAI: {str(e)}")

def add_page_number(doc):
    """Add page numbers in the format 'Page X of Y' to the document footer."""
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add "Page X of Y" field code
        field_code = paragraph.add_run()
        field_code.font.name = 'Calibri Light'
        field_code.font.size = Pt(9)
        
        # Create "Page X of Y" field
        run = field_code._r
        fld_char1 = OxmlElement('w:fldChar')
        fld_char1.set(qn('w:fldCharType'), 'begin')
        run.append(fld_char1)
        
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = ' PAGE \* MERGEFORMAT '
        run.append(instr_text)
        
        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(qn('w:fldCharType'), 'separate')
        run.append(fld_char2)
        
        paragraph.add_run(" de ")
        
        run2 = paragraph.add_run()._r
        fld_char3 = OxmlElement('w:fldChar')
        fld_char3.set(qn('w:fldCharType'), 'begin')
        run2.append(fld_char3)
        
        instr_text2 = OxmlElement('w:instrText')
        instr_text2.set(qn('xml:space'), 'preserve')
        instr_text2.text = ' NUMPAGES \* MERGEFORMAT '
        run2.append(instr_text2)
        
        fld_char4 = OxmlElement('w:fldChar')
        fld_char4.set(qn('w:fldCharType'), 'separate')
        run2.append(fld_char4)
        
        fld_char5 = OxmlElement('w:fldChar')
        fld_char5.set(qn('w:fldCharType'), 'end')
        run2.append(fld_char5)
        
        # End the PAGE field
        run3 = paragraph.add_run()._r
        fld_char6 = OxmlElement('w:fldChar')
        fld_char6.set(qn('w:fldCharType'), 'end')
        run3.append(fld_char6)

def create_formatted_docx(dialogue_data, person_name="XXXX", file_number="CG/0X/2025", date_time="XXXX", committee_members="xxxx y xxxx") -> Optional[bytes]:
    """
    Create a formatted Word document based on the dialogue data.
    
    Args:
        dialogue_data: The structured dialogue data with speakers and text
        person_name: Name of the affected person (default: "XXXX")
        file_number: File number/expediente (default: "CG/0X/2025")
        date_time: Date and time of the interview (default: "XXXX")
        committee_members: Names of committee members (default: "xxxx y xxxx")
    
    Returns:
        bytes: The Word document as bytes or None if there was an error
    """
    try:
        # Parse the JSON string if it's a string
        if isinstance(dialogue_data, str):
            try:
                dialogue_data = json.loads(dialogue_data)
            except:
                return None
        
        # Create a new Word document
        doc = docx.Document()
        
        # Set default font to Calibri Light 11pt
        style = doc.styles['Normal']
        style.font.name = 'Calibri Light'
        style.font.size = Pt(11)
        
        # Set page margins (2.5 cm all around)
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
        
        # Add title with center alignment
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("Comité de Atención de la Violencia de Género")
        title_run.bold = True
        title_run.font.name = 'Calibri Light'
        title_run.font.size = Pt(11)
        
        # Add subtitle with center alignment
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(f"Entrevista a la persona agraviada {person_name}")
        subtitle_run.bold = True
        subtitle_run.font.name = 'Calibri Light'
        subtitle_run.font.size = Pt(11)
        
        # Add expediente with center alignment
        expediente = doc.add_paragraph()
        expediente.alignment = WD_ALIGN_PARAGRAPH.CENTER
        expediente_run = expediente.add_run(f"en el expediente {file_number}¹")
        expediente_run.bold = True
        expediente_run.font.name = 'Calibri Light'
        expediente_run.font.size = Pt(11)
        
        # Add empty line
        doc.add_paragraph()
        
        # Add the footnote text at the bottom of the first page
        footnote_paragraph = doc.add_paragraph()
        footnote_paragraph.paragraph_format.left_indent = Cm(1)
        footnote_paragraph.paragraph_format.first_line_indent = Cm(-0.5)
        footnote_run = footnote_paragraph.add_run("¹ Todas las transcripciones de las entrevistas se realizan eliminando muletillas, frases aisladas, así como frases repetidas.")
        footnote_run.font.name = 'Calibri Light'
        footnote_run.font.size = Pt(8)
        
        # Add introductory text
        intro = doc.add_paragraph()
        intro_text = f"A continuación, se presenta la transcripción íntegra de la entrevista realizada por la Unidad de Trabajo del Comité de Atención de la Violencia de Género ('Comité') dentro del expediente {file_number}, a la persona agraviada en el expediente citado, a quienes se hará referencia en la transcripción con tales términos –y no con su nombres– según corresponda, para identificar sus intervenciones en la entrevista. Los párrafos correspondientes a las intervenciones de la Unidad de Trabajo del Comité, para su clara diferenciación, se presentan en cursivas y separadas con un espacio de las intervenciones correspondientes a la persona agraviada."
        intro.add_run(intro_text)
        intro.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # Add empty line
        doc.add_paragraph()
        
        # Format date and time for the header
        date_header = date_time
        if date_time and date_time != "XXXX":
            # Try to parse and format the date more formally if possible
            try:
                # If date is in DD/MM/YYYY HH:MM format, try to convert it to text format
                parts = date_time.split()
                if len(parts) >= 2:  # Has date and time
                    date_parts = parts[0].split('/')
                    time_parts = parts[1].split(':')
                    
                    if len(date_parts) == 3 and len(time_parts) >= 2:
                        day = int(date_parts[0])
                        month = int(date_parts[1])
                        year = int(date_parts[2])
                        hour = int(time_parts[0])
                        minute = int(time_parts[1])
                        
                        # Spanish month names
                        month_names = [
                            "enero", "febrero", "marzo", "abril", "mayo", "junio",
                            "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
                        ]
                        
                        # Convert to text format
                        date_header = f"las {hour} horas con {minute} minutos del {day} de {month_names[month-1]} del año {year}"
            except:
                # If parsing fails, use the original string
                date_header = date_time
        
        # Add the main interview header with date and time
        header = doc.add_paragraph()
        if date_header == "XXXX":
            header_text = f"Siendo las nueve horas con doce minutos del trece de febrero del año dos mil veinticinco, la Unidad de Trabajo del Comité de Atención de la Violencia de Género, conformada por {committee_members}, actuando dentro del expediente {file_number}, realiza la entrevista a {person_name} en su calidad de persona agraviada."
        else:
            header_text = f"Siendo {date_header}, la Unidad de Trabajo del Comité de Atención de la Violencia de Género, conformada por {committee_members}, actuando dentro del expediente {file_number}, realiza la entrevista a {person_name} en su calidad de persona agraviada."
        header.add_run(header_text).bold = True
        
        # Add empty line
        doc.add_paragraph()
        
        # Add the dialogue
        for entry in dialogue_data.get("dialogue", []):
            speaker = entry.get("speaker", "")
            text = entry.get("text", "")
            
            # Skip empty entries
            if not text.strip():
                continue
            
            # Check if this is the person designated as "Persona Agraviada" by OpenAI
            is_affected_person = "persona agraviada" in speaker.lower() or "agraviada" in speaker.lower()
            
            # Add appropriate formatting based on who is speaking
            if is_affected_person:
                # Persona agraviada format
                paragraph = doc.add_paragraph()
                speaker_run = paragraph.add_run("Persona agraviada: ")
                speaker_run.bold = True
                paragraph.add_run(text)
            else:
                # Interviewer format (italic)
                paragraph = doc.add_paragraph()
                text_run = paragraph.add_run(f"{text}")
                text_run.italic = True
            
            # Add a small space between entries
            doc.add_paragraph()
        
        # Add page numbers in the footer
        add_page_number(doc)
        
        # Save document to a BytesIO object
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return docx_bytes.getvalue()
        
    except Exception as e:
        print(f"Error creating Word document: {str(e)}")
        return None 